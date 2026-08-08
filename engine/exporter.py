"""
engine/exporter.py
Exports a finished transcript to TXT, DOCX, or PDF. Kept separate from
the transcription engine so export logic can evolve independently.
"""

import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics

import config


class ExportError(Exception):
    pass


class Exporter:
    """Writes transcript text to disk in the requested format."""

    @staticmethod
    def _resolve_path(suggested_name: str, extension: str, out_dir: str) -> str:
        os.makedirs(out_dir, exist_ok=True)
        name = suggested_name.strip() or "transcript"
        if not name.lower().endswith(extension):
            name = f"{name}{extension}"
        return os.path.join(out_dir, name)

    # -- TXT -----------------------------------------------------------
    @staticmethod
    def export_txt(text: str, suggested_name: str, out_dir: str = config.OUTPUTS_DIR) -> str:
        path = Exporter._resolve_path(suggested_name, ".txt", out_dir)
        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(text)
        except OSError as exc:
            raise ExportError(f"Could not write TXT file: {exc}") from exc
        return path

    # -- DOCX ------------------------------------------------------------
    @staticmethod
    def export_docx(
        text: str,
        suggested_name: str,
        out_dir: str = config.OUTPUTS_DIR,
        source_file: str = "",
        language: str = "",
    ) -> str:
        path = Exporter._resolve_path(suggested_name, ".docx", out_dir)
        try:
            doc = Document()

            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run(config.APP_NAME)
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)

            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = subtitle.add_run("Transcription Report")
            sub_run.italic = True
            sub_run.font.size = Pt(11)
            sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

            doc.add_paragraph()

            meta = doc.add_paragraph()
            meta_lines = [
                f"Source file: {source_file}" if source_file else None,
                f"Detected language: {language}" if language else None,
                f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            ]
            meta.add_run("\n".join(l for l in meta_lines if l)).font.size = Pt(9)

            doc.add_paragraph()
            body = doc.add_paragraph(text)
            for run in body.runs:
                run.font.size = Pt(11)

            doc.save(path)
        except Exception as exc:
            raise ExportError(f"Could not write DOCX file: {exc}") from exc
        return path

    # -- PDF ---------------------------------------------------------------
    @staticmethod
    def export_pdf(
        text: str,
        suggested_name: str,
        out_dir: str = config.OUTPUTS_DIR,
        source_file: str = "",
        language: str = "",
    ) -> str:
        path = Exporter._resolve_path(suggested_name, ".pdf", out_dir)
        try:
            doc = SimpleDocTemplate(
                path, pagesize=A4,
                leftMargin=2 * cm, rightMargin=2 * cm,
                topMargin=2 * cm, bottomMargin=2 * cm,
            )
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                "TitleBlue", parent=styles["Title"],
                textColor="#3B82F6", fontSize=20,
            )
            meta_style = ParagraphStyle(
                "Meta", parent=styles["Normal"],
                textColor="#64748B", fontSize=9,
            )
            body_style = ParagraphStyle(
                "Body", parent=styles["Normal"],
                fontSize=11, leading=16,
            )

            story = [Paragraph(config.APP_NAME, title_style)]
            story.append(Paragraph("Transcription Report", styles["Italic"]))
            story.append(Spacer(1, 12))

            meta_lines = [
                f"Source file: {source_file}" if source_file else None,
                f"Detected language: {language}" if language else None,
                f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            ]
            for line in filter(None, meta_lines):
                story.append(Paragraph(line, meta_style))
            story.append(Spacer(1, 16))

            # Escape basic XML-sensitive characters for reportlab's mini-markup.
            safe_text = (
                text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            )
            for para in safe_text.split("\n"):
                if para.strip():
                    story.append(Paragraph(para, body_style))
                    story.append(Spacer(1, 8))

            doc.build(story)
        except Exception as exc:
            raise ExportError(f"Could not write PDF file: {exc}") from exc
        return path
